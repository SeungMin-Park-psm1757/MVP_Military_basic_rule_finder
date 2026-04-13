from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import re
from typing import Any

from docx import Document


@dataclass(slots=True)
class ConversationTurn:
    question: dict[str, Any]
    answer: dict[str, Any] | None = None


def build_conversation_turns(history: list[dict[str, Any]]) -> list[ConversationTurn]:
    turns: list[ConversationTurn] = []
    current_question: dict[str, Any] | None = None

    for message in history:
        role = message.get("role")
        if role == "user":
            if current_question is not None:
                turns.append(ConversationTurn(question=current_question))
            current_question = message
            continue

        if role == "assistant":
            if current_question is None:
                turns.append(ConversationTurn(question={"content": ""}, answer=message))
            else:
                turns.append(ConversationTurn(question=current_question, answer=message))
                current_question = None

    if current_question is not None:
        turns.append(ConversationTurn(question=current_question))

    return turns


def _clean_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _append_markdown(document: Document, markdown: str) -> None:
    for raw_line in _clean_text(markdown).splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(line)


def build_conversation_docx(
    history: list[dict[str, Any]],
    *,
    title: str = "군 복무 법규 RAG 대화",
) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph("실무 참고용 안내이며 법률자문이 아닙니다.")

    turns = build_conversation_turns(history)
    for index, turn in enumerate(turns, start=1):
        document.add_heading(f"질문 {index}", level=1)
        document.add_paragraph(_clean_text(str(turn.question.get("content", ""))))

        if turn.answer is None:
            continue

        document.add_heading(f"답변 {index}", level=1)
        _append_markdown(document, str(turn.answer.get("answer_markdown", "")))

        evidence = turn.answer.get("evidence") or []
        if evidence:
            document.add_heading("근거", level=2)
            for hit in evidence:
                chunk = getattr(hit, "chunk", None)
                if chunk is None:
                    continue
                parts = [chunk.law_name, chunk.source_type]
                article_ref = " ".join(part for part in [chunk.article_no, chunk.article_title] if part).strip()
                if article_ref:
                    parts.append(article_ref)
                if chunk.effective_date:
                    parts.append(f"시행일 {chunk.effective_date}")
                if chunk.source_url:
                    parts.append(chunk.source_url)
                document.add_paragraph(" / ".join(parts), style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
